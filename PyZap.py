from zapv2 import ZAPv2
import time
import argparse
from docx import Document
from urllib.parse import urlparse

def parse_args():
    parser = argparse.ArgumentParser(description="ZAP Scanner Script")
    parser.add_argument('target_url', help="Target URL to scan")
    parser.add_argument('-o', '--output', help="Output DOCX file")
    return parser.parse_args()

def create_report(alerts, output_file):
    doc = Document()
    doc.add_heading('ZAP Scan Report', 0)

    for alert in alerts:
        doc.add_heading(alert.get('name', 'No Name'), level=1)
        doc.add_paragraph(alert.get('description', 'No Description'))

    doc.save(output_file)

def main():
    args = parse_args()

    # Set your ZAP API key
    api_key = 'rku4h051aabt31khvmd6ujen0b'

    # Instantiate ZAP object
    zap = ZAPv2(apikey=api_key, proxies={'http': 'http://localhost:8080', 'https': 'http://localhost:8080'})

    # Start a new ZAP session
    zap.core.new_session('example_session', overwrite=True)

    # Set the target URL from the command-line argument
    target_url = args.target_url

    # Start spidering
    zap.spider.scan(target_url)

    # Wait for spider to finish
    while True:
        spider_status = zap.spider.status(target_url)
        print(f"Spider status: {spider_status}%")
        if spider_status == '100':
            break
        time.sleep(1)

    # Start active scanning
    zap.ascan.scan(target_url)

    # Wait for active scan to finish
    while True:
        ascan_status = zap.ascan.status(target_url)
        print(f"Active Scan status: {ascan_status}%")
        if ascan_status == '100':
            break
        time.sleep(1)

    # Fetch and display vulnerabilities (collecting alerts as dictionaries in a list)
    alerts = []
    for alert in zap.core.alerts(baseurl=target_url):
        alert_dict = {'name': alert['name'], 'description': alert['description']}
        if alert_dict not in alerts:
            alerts.append(alert_dict)

    # Determine the output file name based on the entered URL
    if args.output:
        output_file = args.output
    else:
        parsed_url = urlparse(target_url)
        output_file = f"{parsed_url.netloc}.docx"

    # Create a DOCX report
    create_report(alerts, output_file)
    print(f"Report saved to: {output_file}")

if __name__ == "__main__":
    main()

