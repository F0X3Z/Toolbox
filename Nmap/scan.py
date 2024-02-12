"""Nmap scanner which gives you a list with alive hosts"""
from datetime import datetime
import nmap
from docx import Document


def discover_hosts():
    """Start Nmap scan"""
    nm = nmap.PortScanner()
    network_range = '192.168.1.0/24'  # Change this to your actual network range
    nm.scan(hosts=network_range, arguments='-n -sP')

    hosts_list = []
    for host in nm.all_hosts():
        hosts_list.append(host)

    return network_range, hosts_list

def export_to_docx(network_range, hosts_list):
    document = Document()
    current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    document.add_heading('Hosts on {}'.format(network_range), level=1)
    document.add_paragraph('Scan performed at: {}'.format(current_time))
    document.add_paragraph()

    for host in hosts_list:
        document.add_paragraph(host)
        print(host)  # Print the host as well

    filename = 'Hosts_on_{}.docx'.format(network_range.replace('/', '_'))
    document.save(filename)
    print('Exported hosts to', filename)

if __name__ == "__main__":
    network_range, hosts = discover_hosts()
    export_to_docx(network_range, hosts)
